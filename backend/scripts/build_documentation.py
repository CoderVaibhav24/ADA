"""Regenerate ADA_Vision_Project_Documentation.docx from the as-built system.

The original document described a single-model plan (one fine-tuned foundation
model). The delivered system is a multi-model pipeline, so this script rebuilds
the document to match what actually runs, and drops the sections that described
components never built (STAC cataloguing, Celery/Redis, TerraTorch fine-tuning
as the core detector, learning-phase timeline).

Run:  .venv\\Scripts\\python.exe scripts\\build_documentation.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[2] / "ADA_Vision_Project_Documentation.docx"

doc = Document()


# --------------------------------------------------------------- helpers
def title(text: str) -> None:
    doc.add_paragraph(text, style="Title")


def h1(text: str) -> None:
    doc.add_heading(text, level=1)


def h2(text: str) -> None:
    doc.add_heading(text, level=2)


def para(text: str, *, bold: bool = False, italic: bool = False,
         center: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def bullet(text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def table(headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for cell, head in zip(t.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    doc.add_paragraph()


def decision(name: str, rationale: str) -> None:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    doc.add_paragraph(rationale)


# --------------------------------------------------------------- cover
title("Project ADA-Vision")
para("AI-Powered Bi-Temporal Land Change Detection System "
     "for Agra Development Authority (ADA)", bold=True, center=True)
para("Technical Design & Project Proposal Document\n"
     "Prepared by: Vaibhav Sahay\nDate: August 2026", center=True)

h1("Table of Contents")
for line in [
    "1. Executive Summary",
    "2. Problem Statement — Why This Matters",
    "3. Why This Is a Critical Priority for ADA",
    "4. Proposed Solution Approach",
    "5. System Architecture & Design Decisions",
    "6. The Multi-Model Detection Ensemble",
    "7. Technical Requirements & Tech Stack",
    "8. Data Pipeline & Workflow",
    "9. File Formats & Data Standards",
    "10. Model Strategy: AI Mode vs Diff Mode",
    "11. Licensing Strategy",
    "12. Continuous Learning & Feedback Loop",
    "13. Accuracy: Current Position and the Route to Production Grade",
    "14. Tooling: AI-Assisted Development & Compute",
    "15. Timeline — Solo Developer Execution Plan",
    "16. Risk Factors & Continuity Planning",
    "17. Budget & Resource Requirements",
    "18. Success Metrics",
]:
    doc.add_paragraph(line)

# --------------------------------------------------------------- 1
h1("1. Executive Summary")
para(
    "The Agra Development Authority (ADA) requires an AI-based system capable of "
    "comparing satellite and drone imagery captured at two different points in time "
    "for the same geographic area, and automatically identifying unauthorized changes "
    "on the ground. This includes illegal construction, encroachment on public and "
    "government land, and development within designated red zones or no-construction "
    "zones. This document lays out the technical design, architecture, model strategy, "
    "licensing position, and execution plan for delivering this system as an "
    "enterprise-grade product."
)
para(
    "The system ingests two time-stamped images of the same location "
    "(satellite-to-satellite, drone-to-drone, or cross-source satellite-to-drone), "
    "aligns them onto a common grid, detects physical changes, classifies the nature "
    "of each change, and cross-references detected changes against zoning boundary "
    "data to flag violations. Officers review flagged violations through a web "
    "dashboard, and their corrections are captured as labelled data that feeds "
    "periodic model retraining."
)
para(
    "A key design point, and the main revision in this version of the document: "
    "detection is not performed by a single model. It is performed by an ensemble of "
    "several specialised models and deterministic stages chained together, each "
    "solving one part of the problem that a single end-to-end network solves poorly. "
    "Section 6 describes this ensemble in full."
)

# --------------------------------------------------------------- 2
h1("2. Problem Statement — Why This Matters")
para(
    "ADA currently relies on manual site inspections and citizen complaints to "
    "identify illegal construction, encroachment, and red-zone violations. This "
    "approach is slow, inconsistent, resource-intensive, and reactive rather than "
    "proactive. By the time a violation is discovered, construction is often complete, "
    "making enforcement legally and practically difficult."
)
para("The core problem to solve is:")
for b in [
    "Given two images of the same location captured at different times, automatically "
    "identify what has physically changed on the ground.",
    "Classify each change: new construction, structural expansion, vegetation or land "
    "clearance, demolition.",
    "Cross-reference each detected change against ADA's zoning master plan "
    "(residential, commercial, red zone, no-construction zone, riverbank buffer "
    "zones, and similar).",
    "Flag violations automatically with supporting visual evidence, location "
    "coordinates, and a confidence score.",
    "Handle images from mixed sources — satellite imagery and drone footage — which "
    "differ in resolution, viewing angle, and colour characteristics.",
]:
    bullet(b)

# --------------------------------------------------------------- 3
h1("3. Why This Is a Critical Priority for ADA")
for b in [
    "Scale problem: Manual inspection cannot cover the entire ADA jurisdiction "
    "regularly; automated monitoring can scan the full area on every imagery refresh "
    "cycle.",
    "Legal enforcement window: Early detection during construction (foundation or "
    "structure stage) gives ADA legal standing to act before a structure is completed "
    "and occupied.",
    "Revenue and land protection: Red-zone and riverbank encroachment often causes "
    "long-term environmental and flood-risk damage that is costly to reverse after "
    "the fact.",
    "Audit trail and accountability: An automated system creates a timestamped, "
    "evidence-backed record of violations, reducing disputes and the corruption risk "
    "associated with manual reporting.",
    "Precedent for other authorities: A working solution for ADA can be replicated "
    "for other development authorities and municipal bodies, creating reuse value "
    "beyond this single engagement.",
]:
    bullet(b)

# --------------------------------------------------------------- 4
h1("4. Proposed Solution Approach")
para(
    "The system is a staged pipeline that combines deep learning for visual "
    "understanding with deterministic geospatial rule-checking for zone violations. "
    "This hybrid approach (AI plus rules engine) ensures that zoning decisions remain "
    "transparent, auditable, and legally defensible — violations are flagged based on "
    "explicit polygon geometry rules, not opaque model judgment alone."
)
h2("4.1 Pipeline Stages")
for b in [
    "Stage 1 — Preprocessing and co-registration: reproject both epochs onto one "
    "common working grid, correct residual misalignment to sub-pixel accuracy, and "
    "normalise colour so that sensor differences are not mistaken for change.",
    "Stage 2 — Structure extraction: segment building footprints independently in "
    "each epoch, rather than asking one network to compare two images directly.",
    "Stage 3 — Change reasoning: diff the two footprint maps and confirm each "
    "candidate against independent evidence (loss of vegetation, colour change) "
    "before it is allowed to become a detection.",
    "Stage 4 — Full-structure refinement: expand each confirmed detection to the "
    "complete building outline so the officer sees a whole structure, not a fragment.",
    "Stage 5 — Zone violation logic: convert detections to spatial polygons and test "
    "them against ADA's zoning boundaries using deterministic geometric rules.",
    "Stage 6 — Officer review: every detection is adjudicated by a human before it "
    "becomes an official record, and that adjudication is captured as training data.",
]:
    bullet(b)

h2("4.2 Why Not a Single End-to-End Model")
para(
    "The obvious design — one bi-temporal change-detection network that takes both "
    "images and outputs a change mask — was built first and evaluated on real Agra "
    "data. It underperformed for a specific and instructive reason: the two epochs "
    "come from different sensors. A satellite image and a drone image of the same "
    "unchanged building differ in colour, resolution, and viewing angle, and an "
    "end-to-end comparator learns to fire on exactly those differences. It flagged "
    "vegetation boundaries and roof edges while missing real construction."
)
para(
    "Segmenting each epoch independently avoids this failure mode entirely. A "
    "building-footprint model asked \"is this a building?\" gives an answer that does "
    "not depend on what the other epoch looked like, so a colour or angle difference "
    "between sensors cannot by itself produce a detection. This is the single most "
    "important architectural decision in the system, and it is the reason the design "
    "is an ensemble rather than one model."
)

# --------------------------------------------------------------- 5
h1("5. System Architecture & Design Decisions")
h2("5.1 High-Level Architecture")
para(
    "The system follows a modular architecture: a React frontend for visualisation, a "
    "FastAPI backend for orchestration and business logic, PostgreSQL for relational "
    "and spatial data, and a model-serving layer for inference. This separation "
    "ensures the ML pipeline can be scaled, retrained, or swapped independently of "
    "the application layer."
)
table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Presentation Layer", "React + TypeScript + MapLibre GL",
         "Map-based UI for officers to view flagged violations overlaid on imagery "
         "and zoning boundaries, with an integrated review queue"],
        ["API / Orchestration Layer", "FastAPI",
         "Handles requests, queues inference jobs, serves results and map tiles, "
         "captures officer feedback"],
        ["Authentication Layer", "SuperTokens (self-hosted)",
         "Session-based authentication; no user credentials leave ADA infrastructure"],
        ["Inference Layer", "PyTorch (GPU) + ONNX Runtime (CPU), in-process workers",
         "Runs the model ensemble on submitted image pairs"],
        ["Data Layer", "PostgreSQL + SQLAlchemy",
         "Stores structured metadata, zoning polygons, detection geometry, and "
         "officer review records"],
        ["Storage Layer", "File-backed raster store (object-storage compatible)",
         "Stores large GeoTIFF/COG imagery and generated masks outside the database"],
        ["Tile Server", "rio-tiler / titiler",
         "Serves imagery and change-heat overlays as web map tiles on demand"],
        ["Job Runner", "Background worker queue",
         "Runs long inference jobs asynchronously with progress reporting, so the "
         "API stays responsive"],
    ],
)

h2("5.2 Key Architecture Decisions")
decision(
    "Decision 1: Hybrid AI + rules engine, not pure end-to-end AI",
    "Rationale: Zoning violations are ultimately a legal and geometric determination "
    "— is this polygon inside a red zone? Encoding that as a deterministic spatial "
    "query rather than letting a neural network decide keeps the system explainable "
    "and defensible in enforcement proceedings. The models detect and describe "
    "physical change; the rules engine determines legal violation status.",
)
decision(
    "Decision 2: An ensemble of specialised models, not one general model",
    "Rationale: The problem decomposes into sub-problems with different failure "
    "modes — alignment, structure detection, change confirmation, outline recovery. "
    "Each is solved by a component chosen for that job, and each can be replaced or "
    "retrained independently without rebuilding the pipeline. This also means a "
    "weakness in one stage is visible and fixable rather than hidden inside a single "
    "opaque network. Section 6 details the ensemble.",
)
decision(
    "Decision 3: Segment each epoch independently, then diff the results",
    "Rationale: This is what makes cross-sensor comparison viable. Because each "
    "epoch is judged on its own, a colour cast, a resolution difference, or an "
    "off-nadir drone angle cannot by itself generate a change detection. The trade-off "
    "is that segmentation errors in either epoch propagate into the diff, which is "
    "why a candidate detection must be independently confirmed before it is reported.",
)
decision(
    "Decision 4: Confirmation gating instead of raw differencing",
    "Rationale: The dominant false positive is a building the segmenter found in the "
    "second epoch but simply missed in the first — nothing changed on the ground. "
    "Every candidate must therefore be corroborated by independent evidence: the area "
    "was vegetation before and is built now, or the imagery itself demonstrably "
    "changed there. On real Agra data this gating reduced false detection area by "
    "roughly two orders of magnitude.",
)
decision(
    "Decision 5: Raster-to-vector conversion for violation storage",
    "Rationale: Model outputs are raster pixel masks, but zoning boundaries and legal "
    "records are vector polygons. Converting detections into vector polygons "
    "immediately after inference allows consistent spatial querying, area "
    "calculation, and reporting, and makes each detection an addressable record that "
    "an officer can adjudicate.",
)
decision(
    "Decision 6: Asynchronous inference with progress reporting",
    "Rationale: Inference over a full orthophoto takes minutes. Running it inside an "
    "API request would block the server and time out the browser. Jobs are queued and "
    "execute in the background, writing stage and progress into the database, which "
    "the dashboard polls — so the officer sees which model is currently running "
    "rather than an unexplained wait.",
)
decision(
    "Decision 7: Separate Diff Mode and AI Mode as distinct execution paths",
    "Rationale: Not every comparison needs the full ensemble. A lightweight "
    "difference mode gives officers an instant visual preview for manual screening, "
    "while AI Mode is reserved for cases requiring an official, evidence-backed "
    "report. This balances speed and compute cost against accuracy needs.",
)

# --------------------------------------------------------------- 6
h1("6. The Multi-Model Detection Ensemble")
para(
    "Detection is performed by several models and deterministic stages chained "
    "together. Each row below is a distinct component with a distinct job. All are "
    "permissively licensed and run on infrastructure ADA controls."
)
table(
    ["#", "Component", "Job in the pipeline", "Where it runs"],
    [
        ["1", "FFT phase cross-correlation (classical)",
         "Sub-pixel co-registration of the two epochs after reprojection onto a "
         "common grid. Recovers residual translation between a drone flight and a "
         "satellite pass; rejects implausible shifts rather than trusting them.",
         "CPU"],
        ["2", "Radiometric normalisation and false-colour handling (classical)",
         "Histogram matching between epochs, plus automatic detection and conversion "
         "of colour-infrared composites, so a false-colour satellite epoch can be "
         "compared against a true-colour drone epoch.",
         "CPU"],
        ["3", "Building-footprint segmentation network (U-Net, ONNX)",
         "Run separately on each epoch to answer 'where are the buildings?' — the "
         "structural backbone of the system. Tiled across the scene in overlapping "
         "chips with smooth blending so large orthophotos are processed at full "
         "resolution without visible seams.",
         "CPU (ONNX Runtime)"],
        ["4", "Vegetation indices (NDVI / excess-green)",
         "Computed per epoch from raw bands. Provides the 'was green, is now built' "
         "evidence that confirms encroachment on green space without needing a pixel "
         "comparison — the ADA core case.",
         "CPU"],
        ["5", "Colour and structure difference operators",
         "Colour difference catches bare-ground-to-building change; a "
         "colour-invariant gradient/edge difference provides a comparison that "
         "survives satellite-to-drone colour shifts. Used as confirmation evidence in "
         "AI Mode and as the entire detector in Diff Mode.",
         "CPU"],
        ["6", "SAM 2 (Segment Anything Model 2)",
         "Full-structure refinement. Once a detection is confirmed, SAM 2 is prompted "
         "with its bounding box on the later epoch and returns the complete building "
         "outline, so officers see whole structures rather than fragments. Robust to "
         "shadow, low light, and off-nadir viewing angle.",
         "GPU"],
        ["7", "Bi-temporal change-detection backend (alternate path)",
         "A general change-detection route retained for scenes where structure "
         "segmentation is not the right tool — deep feature-distance comparison, with "
         "a drop-in slot for a purpose-built change-detection network.",
         "GPU / CPU"],
        ["8", "Deterministic zone rules engine",
         "Not a model. Vectorises the detection mask, computes true geodesic areas, "
         "and tests each polygon against ADA zoning geometry to assign violation "
         "status. This is the component that decides the word 'illegal'.",
         "CPU"],
    ],
)
para(
    "Every analysis run records which components executed, and that list is shown in "
    "the dashboard and included in the exported report. An officer or auditor can "
    "therefore see exactly how a given violation was produced — a requirement for "
    "evidence used in enforcement.",
    italic=True,
)

h2("6.1 How the Components Combine")
para(
    "The ensemble is a chain with a confirmation gate in the middle, not a vote. "
    "Components 1 and 2 produce a comparable pair. Component 3 produces a footprint "
    "map for each epoch, and the difference between those maps produces candidate "
    "detections. Components 4 and 5 act as the gate: a candidate is only promoted to "
    "a detection if independent evidence supports it. Component 6 then completes the "
    "geometry of whatever survives, and component 8 decides its legal status."
)
para(
    "The consequence worth stating plainly for reviewers: the system is designed to "
    "under-report rather than over-report. A missed structure costs one inspection; a "
    "false accusation of illegal construction costs ADA credibility and creates legal "
    "exposure. The gating is tuned accordingly, and it is adjustable per deployment."
)

h2("6.2 Where the Models Live — Offline and Air-Gapped Operation")
para(
    "All model weights are vendored into the deployment itself rather than fetched "
    "from the internet at runtime. A single setup command downloads every model into "
    "the application's own weights directory and records the exact upstream commit of "
    "each one in a manifest file; from that point the system loads entirely from "
    "local disk and requires no network access to perform detection."
)
para("This matters for three reasons a government deployment cares about:")
for b in [
    "Air-gapped operation: an ADA server with no internet access runs the full "
    "pipeline. Deployment is a file copy, not a download.",
    "Reproducibility: the manifest pins each model to a specific commit hash, so the "
    "system deployed a year from now behaves identically to the one demonstrated "
    "today, regardless of what changes upstream.",
    "Data sovereignty: no imagery, detection, or officer decision is sent to any "
    "third-party AI service at any point. Every model runs on hardware ADA controls, "
    "which also means there is no per-inference or per-scene cost.",
]:
    bullet(b)
para(
    "Total vendored model size is approximately 261 MB across the ensemble — small "
    "enough to ship with the application and to hold in memory on modest hardware.",
    italic=True,
)

# --------------------------------------------------------------- 7
h1("7. Technical Requirements & Tech Stack")
table(
    ["Component", "Chosen Technology"],
    [
        ["Backend framework", "FastAPI"],
        ["Frontend framework", "React + TypeScript (Vite)"],
        ["ORM", "SQLAlchemy"],
        ["Database", "PostgreSQL"],
        ["Authentication", "SuperTokens (self-hosted)"],
        ["Deep learning framework", "PyTorch (CUDA build for GPU inference)"],
        ["Optimised model runtime", "ONNX Runtime (CPU inference for the segmenter)"],
        ["Model distribution", "Hugging Face Hub (automatic weight download and caching)"],
        ["Foundation segmentation model", "SAM 2 via Transformers"],
        ["Classical CV", "OpenCV, scikit-image, SciPy"],
        ["Geospatial raster handling", "Rasterio / GDAL, rio-cogeo"],
        ["Geospatial vector handling", "Shapely, PyProj"],
        ["Map tile serving", "rio-tiler / titiler"],
        ["Map rendering", "MapLibre GL JS"],
        ["Containerisation", "Docker (database and auth services)"],
        ["Training compute environment", "Local NVIDIA GPU, with Google Colab for larger fine-tuning runs"],
        ["Version control", "Git / GitHub"],
    ],
)

# --------------------------------------------------------------- 8
h1("8. Data Pipeline & Workflow")
h2("8.1 End-to-End Workflow")
for n in [
    "Image ingestion: Two time-stamped images (satellite or drone) of the same area "
    "are uploaded through the dashboard. Each is converted to a Cloud-Optimized "
    "GeoTIFF, its coordinate reference system and ground resolution are recorded, and "
    "it becomes available as a map layer.",
    "Superimposing: Both epochs are reprojected onto a single common working grid, "
    "aligned to sub-pixel accuracy by phase correlation, and radiometrically "
    "normalised so that sensor and lighting differences are not read as change.",
    "Structure extraction: The building-footprint network is run independently over "
    "each epoch, tiled in overlapping chips and blended, producing one footprint "
    "probability map per epoch.",
    "Change reasoning: The two footprint maps are differenced to produce candidate "
    "new construction, and each candidate is tested against independent evidence "
    "(vegetation loss, colour change) before being accepted.",
    "Full-structure refinement: SAM 2 expands each accepted detection to the complete "
    "building outline on the later epoch.",
    "Mask generation: The final probability map is written as a georeferenced "
    "change-heat COG and served to the map as an overlay layer.",
    "Vectorisation and classification: The mask is converted into vector polygons; "
    "speckle below the minimum area threshold is discarded; each polygon receives a "
    "change-type label, a true geodesic area, and a confidence score.",
    "Zone violation check: Each polygon is tested against the project's zoning "
    "geometry. Intersecting polygons are flagged as illegal encroachment, with the "
    "percentage of overlap recorded.",
    "Officer review: Officers review each detection in the dashboard with a "
    "before/after evidence patch, and mark it confirmed or a false positive.",
    "Reporting and feedback: Confirmed violations are exported as a CSV violation "
    "register and a GeoJSON evidence pack. All adjudications accumulate as labelled "
    "data for the next retraining cycle.",
]:
    numbered(n)

# --------------------------------------------------------------- 9
h1("9. File Formats & Data Standards")
table(
    ["Format", "Usage"],
    [
        ["GeoTIFF / Cloud-Optimized GeoTIFF (COG)",
         "Primary format for all raster imagery, both satellite and drone. Preserves "
         "full pixel fidelity with lossless compression and supports partial, tiled "
         "reads for efficient tile serving."],
        ["World file (.tfw) + projection file (.prj) sidecars",
         "Accepted alongside a plain TIFF when the georeferencing is not embedded — "
         "common with drone processing software exports."],
        ["GeoJSON",
         "Zoning boundaries, detected violation polygons, API responses, and the "
         "exported evidence pack and training-data set."],
        ["CSV",
         "Tabular violation register for enforcement officers, with coordinates, "
         "area, confidence, zone overlap, and review status per detection."],
        ["PNG (intermediate)",
         "Map tiles and before/after evidence patches only; never a source of truth "
         "for georeferenced data."],
    ],
)
para(
    "Note: Proprietary lossy raster formats such as ECW are explicitly avoided for "
    "any data used in model inference or training, as their compression artefacts "
    "degrade change-detection accuracy at exactly the building edges the system "
    "depends on. GeoTIFF/COG is the mandated standard for all raster storage and "
    "processing."
)

# --------------------------------------------------------------- 10
h1("10. Model Strategy: AI Mode vs Diff Mode")
para(
    "Both modes are available per analysis run and are selected by the officer when "
    "starting a comparison. The mode used is recorded against the run and appears on "
    "every exported report, so an evidence-grade result can never be confused with a "
    "triage preview."
)
table(
    ["Mode", "What it runs", "Typical latency", "Use case"],
    [
        ["Diff Mode",
         "Co-registration and normalisation, then classical colour and "
         "colour-invariant structure difference, with vegetation-only change "
         "suppressed. No neural inference.",
         "Seconds",
         "Quick visual screening and manual triage across many locations before "
         "committing to full analysis"],
        ["AI Mode",
         "The full ensemble: co-registration, per-epoch building segmentation, "
         "confirmation gating, SAM 2 full-structure refinement, vectorisation and "
         "zone-violation checking.",
         "Minutes (asynchronous)",
         "Official violation reports with evidence, used for enforcement action"],
    ],
)
para(
    "Diff Mode is deliberately biased toward recall — it is better for a triage pass "
    "to show an officer something that turns out to be nothing than to hide it. AI "
    "Mode is biased toward precision, because its output carries enforcement weight."
)

# --------------------------------------------------------------- 11
h1("11. Licensing Strategy")
para(
    "As this system will be delivered to a government authority, all models and "
    "libraries carry permissive open-source licences (Apache 2.0, MIT, BSD) so that "
    "ADA retains full rights to operate, modify, and redistribute the system without "
    "dependency on any external licensing agreement, and with no per-seat or "
    "per-inference cost."
)
table(
    ["Component", "Licence family"],
    [
        ["Building-footprint segmentation model and weights", "Apache 2.0"],
        ["SAM 2 (Segment Anything Model 2)", "Apache 2.0"],
        ["PyTorch, torchvision", "BSD 3-Clause"],
        ["ONNX Runtime", "MIT"],
        ["Transformers, Hugging Face Hub", "Apache 2.0"],
        ["OpenCV", "Apache 2.0"],
        ["scikit-image, SciPy, NumPy", "BSD 3-Clause"],
        ["Rasterio, Shapely, PyProj, rio-tiler", "BSD 3-Clause / MIT"],
        ["FastAPI, SQLAlchemy, PostgreSQL", "MIT / PostgreSQL Licence"],
        ["MapLibre GL JS", "BSD 3-Clause"],
        ["SuperTokens", "Apache 2.0"],
    ],
)
para(
    "A THIRD_PARTY_LICENSES.md file is maintained in the project repository listing "
    "every third-party model and library used, its licence, and the exact version or "
    "commit hash pinned at time of integration. Model weights are pinned by revision "
    "and mirrored locally, so a change or withdrawal upstream cannot affect a "
    "deployed system. This provides full compliance and auditability for government "
    "procurement review."
)

# --------------------------------------------------------------- 12
h1("12. Continuous Learning & Feedback Loop")
para(
    "The system improves over time through a structured human-in-the-loop active "
    "learning cycle, rather than an unconstrained reinforcement learning approach. "
    "This is a deliberate choice: reinforcement learning requires a reward signal "
    "from repeated environment interaction, which does not naturally exist for a "
    "static image analysis task. A supervised feedback loop achieves the same "
    "practical goal — the system getting more accurate over time — in a way that is "
    "stable, predictable, and appropriate for an enforcement-grade system."
)
for n in [
    "Every detection appears in the officer's review queue in the dashboard, with a "
    "before/after evidence patch and its location on the map.",
    "The officer marks each one Confirmed or False Positive. The decision, the "
    "officer's identity, and the timestamp are stored against that detection; the map "
    "immediately reflects the verified picture.",
    "Confirmed and rejected detections accumulate across every run in the project as "
    "positive and negative labelled examples, exportable as a single GeoJSON training "
    "set with one command from the dashboard.",
    "On a scheduled cadence, that verified set is used to fine-tune the "
    "building-footprint model on Agra-specific imagery — the step that closes the "
    "accuracy gap described in Section 13.",
    "Precision, recall, and false-positive rate are tracked release over release, and "
    "an updated model is promoted to production only if it improves on the incumbent "
    "against a held-out labelled test set.",
]:
    numbered(n)
para(
    "Critically, the model is never retrained on its own unverified output. Only "
    "human-adjudicated detections enter the training set, which prevents the model "
    "from reinforcing its own errors over successive cycles.",
    italic=True,
)

# --------------------------------------------------------------- 13
h1("13. Accuracy: Current Position and the Route to Production Grade")
para(
    "This section is included deliberately, because it is the question a technical "
    "reviewer should ask and the answer determines the remaining schedule."
)
para(
    "The pipeline is complete and every stage works end to end on real Agra imagery. "
    "The limiting factor is not the architecture; it is that the building-footprint "
    "model currently runs on general-purpose open weights trained on aerial imagery "
    "from elsewhere. On Agra's dense, low-contrast rooftops it under-detects, and "
    "because two epochs are segmented independently, an inconsistency between them "
    "propagates into the difference. The confirmation gating described in Section 6 "
    "suppresses the resulting false positives effectively, but it cannot recover a "
    "building the model never saw."
)
para("The route to production-grade accuracy is therefore specific and bounded:")
for b in [
    "Label a set of building chips on ADA's own imagery — a few hundred is the "
    "working estimate, drawn from areas ADA already knows well.",
    "Fine-tune the footprint model on that set. The pipeline around it does not "
    "change; only the weights it loads change.",
    "Gate promotion on a measured F1 score against a held-out Agra test set, rather "
    "than on visual inspection.",
    "From then on, the officer review queue supplies labelled data continuously, so "
    "each retraining cycle draws on a larger and more representative set than the "
    "last.",
]:
    bullet(b)
para(
    "Stating this plainly is a deliberate choice. A system of this kind that claims "
    "high accuracy on general-purpose weights, without local labelled data, should "
    "not be believed — and if ADA acts on such claims and the system is wrong, the "
    "cost lands on the authority. The honest position is that the engineering is "
    "done and the localisation step is scheduled and quantified."
)

# --------------------------------------------------------------- 14
h1("14. Tooling: AI-Assisted Development & Compute")
para(
    "Given the technical breadth of this project — backend, ML pipeline, geospatial "
    "processing, and frontend — executed by a single developer within a compressed "
    "timeline, AI-assisted coding tools and GPU compute are necessary "
    "force-multipliers rather than optional conveniences."
)
para("Requirement 1: Claude Max Plan ($100/month subscription)", bold=True)
for b in [
    "Scaffolding and reviewing code across the FastAPI, SQLAlchemy, and geospatial "
    "integration layers.",
    "Debugging model pipeline issues — tensor shape mismatches, coordinate reference "
    "system errors, alignment failures — which are the highest-time-cost defects in "
    "this domain.",
    "Accelerating the fine-tuning and evaluation scripts described in Section 13.",
    "Sustaining enterprise-grade code quality and documentation despite solo "
    "development, which directly supports the continuity plan in Section 16.",
]:
    bullet(b)
para("Requirement 2: Google Colab Pro ($12/month subscription)", bold=True)
for b in [
    "Provides GPU capacity for fine-tuning runs beyond what the local development "
    "GPU can hold, without upfront hardware investment.",
    "Enables faster iteration during model development compared to CPU-only training.",
    "Cost-effective for a time-boxed solo engagement compared to provisioning "
    "dedicated cloud GPU infrastructure.",
    "Used during the fine-tuning phase only; production inference runs on ADA-"
    "provisioned hardware, and no ADA imagery is required to leave that environment "
    "for routine operation.",
]:
    bullet(b)
para(
    "These two tools together form the minimum viable development environment for "
    "this engagement, and their cost is included in Section 17."
)

# --------------------------------------------------------------- 15
h1("15. Timeline — Solo Developer Execution Plan")
para(
    "Total estimated duration for a solo developer to deliver this system end to end, "
    "through to a working production-grade deployment, is 3.5 months."
)
table(
    ["Phase", "Duration", "Deliverables"],
    [
        ["Phase 1", "Weeks 1-3",
         "Platform foundation: project and imagery management, raster ingestion to "
         "COG, coordinate reference system handling, tile serving, authentication, "
         "and the map dashboard."],
        ["Phase 2", "Weeks 4-6",
         "Preprocessing and superimposing: reprojection onto a common grid, "
         "sub-pixel co-registration, radiometric normalisation, and false-colour "
         "handling for cross-sensor pairs."],
        ["Phase 3", "Weeks 7-9",
         "The detection ensemble: per-epoch building segmentation, footprint "
         "differencing, confirmation gating, and SAM 2 full-structure refinement; "
         "both AI Mode and Diff Mode execution paths."],
        ["Phase 4", "Weeks 10-11",
         "Zone-violation rules engine and reporting: vectorisation, geodesic area "
         "computation, zoning intersection logic, violation register and evidence "
         "pack exports."],
        ["Phase 5", "Weeks 12-13",
         "Officer review workflow and the feedback loop: review queue, adjudication "
         "capture, and labelled-dataset export for retraining."],
        ["Phase 6", "Weeks 14-15",
         "Agra-specific fine-tuning against labelled local chips, accuracy "
         "validation against a held-out test set, deployment hardening, and handover "
         "documentation."],
    ],
)
para("Total: approximately 15 weeks (3.5 months), inclusive of development, "
     "model localisation, and integration.")

# --------------------------------------------------------------- 16
h1("16. Risk Factors & Continuity Planning")
h2("16.1 Key Risks")
table(
    ["Risk", "Severity", "Mitigation"],
    [
        ["Single point of failure (solo developer)", "High",
         "Thorough documentation, commented code, and architecture decision records "
         "(this document) so a successor can pick up the project mid-stream."],
        ["Accuracy on general-purpose weights below enforcement grade", "High",
         "Explicitly scheduled as Phase 6 and quantified in Section 13; promotion "
         "gated on a measured F1 score against a held-out Agra test set rather than "
         "visual inspection."],
        ["Cross-sensor domain gap (satellite vs drone)", "Medium",
         "Addressed architecturally by segmenting each epoch independently, plus "
         "radiometric normalisation, false-colour handling, and colour-invariant "
         "confirmation signals."],
        ["Limited labelled ADA-specific training data", "Medium",
         "Fine-tuning a pretrained model needs far less data than training from "
         "scratch; the officer review loop grows the labelled set continuously "
         "during normal operation."],
        ["Off-nadir drone imagery and parallax", "Medium",
         "Sub-pixel co-registration with rejection of implausible shifts, and "
         "SAM 2 refinement which recovers full structure outlines under oblique "
         "viewing angles."],
        ["Compute resource availability for training", "Medium",
         "Local GPU for routine work with Colab Pro for larger runs; production "
         "inference is designed to run on modest hardware."],
        ["False positives or negatives affecting enforcement decisions", "High",
         "Zone-violation determination stays in the deterministic rules engine, not "
         "the model; officer confirmation is required before any flag becomes an "
         "official record; the pipeline is tuned to under-report rather than "
         "over-report."],
    ],
)
h2("16.2 Continuity Plan (If Developer Leaves Mid-Project)")
for b in [
    "All code is maintained in version control with clear commit history and "
    "module-level documentation.",
    "This document is the architecture reference — an incoming developer can resume "
    "from the documented pipeline stage without re-deriving design decisions.",
    "The ensemble is modular by design: each model is loaded behind a small backend "
    "interface, so a successor can replace or retrain one component without "
    "understanding the whole system.",
    "Model weights are pinned by revision and cached locally; training scripts, "
    "configuration, and checkpoints are versioned separately from application code so "
    "training can be reproduced or continued.",
    "Runtime behaviour — which model runs, thresholds, gating mode — is controlled by "
    "configuration rather than code changes, so tuning does not require the original "
    "developer.",
    "Third-party licence documentation ensures legal continuity regardless of who "
    "maintains the system.",
]:
    bullet(b)

# --------------------------------------------------------------- 17
h1("17. Budget & Resource Requirements")
table(
    ["Item", "Cost", "Duration", "Notes"],
    [
        ["Claude Max Plan (AI-assisted development)", "$100/month", "~3.5 months", "~$350"],
        ["Google Colab Pro (GPU compute for fine-tuning)", "$12/month", "~3.5 months", "~$42"],
        ["Raster storage (imagery hosting)", "Variable", "Ongoing",
         "Scales with imagery volume; file-backed and object-storage compatible"],
        ["All models, weights, and libraries", "Free (Apache 2.0 / MIT / BSD)", "N/A", "$0"],
        ["Inference hardware (production)", "One-time / existing", "Ongoing",
         "A single NVIDIA GPU is sufficient; the footprint segmenter runs on CPU"],
    ],
)
para(
    "Estimated total recurring tooling cost for the 3.5-month build phase is "
    "approximately $392 (Claude Max plus Google Colab Pro), excluding variable "
    "storage costs."
)
para(
    "The entire software stack — FastAPI, React, PostgreSQL, PyTorch, ONNX Runtime, "
    "and every model in the ensemble including SAM 2 — is free and open source under "
    "permissive licences. There are no per-inference, per-seat, or per-scene fees, "
    "and no third-party AI service is called at runtime, so ADA imagery never leaves "
    "ADA infrastructure during normal operation. The two recurring costs are "
    "development-phase tooling that reduces delivery time and risk for a solo "
    "engagement, and are considerably lower than provisioning dedicated cloud GPU "
    "servers or adding development headcount."
)

# --------------------------------------------------------------- 18
h1("18. Success Metrics")
for b in [
    "Detection accuracy: precision and recall of the ensemble in correctly "
    "identifying genuine new construction between image pairs, validated against a "
    "held-out labelled Agra test set.",
    "False positive rate: the primary metric for officer trust — every false flag "
    "costs an inspection visit and erodes confidence in the system.",
    "Classification accuracy: correct categorisation of change type (new "
    "construction, clearance, demolition, surface change).",
    "Zone violation flagging accuracy: percentage of flagged violations confirmed as "
    "valid by ADA officers during review.",
    "Officer confirmation rate over time: the share of detections confirmed rather "
    "than rejected, tracked release over release, which is the direct measure of "
    "whether the feedback loop is working.",
    "Processing turnaround: time from image pair submission to violation report in AI "
    "Mode, and to preview in Diff Mode.",
    "System adoption: number of officers actively using the dashboard and volume of "
    "image-pair comparisons processed per month post-launch.",
]:
    bullet(b)

doc.add_paragraph()
para("— End of Document —", center=True, italic=True)

for section in doc.sections:
    section.footer.paragraphs[0].text = "Project ADA-Vision — Technical Design Document"

# normalise body font size a little
style = doc.styles["Normal"]
style.font.size = Pt(10.5)

doc.save(OUT)
print(f"wrote {OUT} ({OUT.stat().st_size:,} bytes)")
